"""Text-to-PDF/DOCX conversion for tailored resumes and cover letters.

Parses the structured text resume format, renders via an HTML/CSS template,
and exports to PDF using headless Chromium via Playwright, or to DOCX using
python-docx.

Supported formats: "pdf" (default), "docx".
"""

import logging
from pathlib import Path

from applypilot.config import TAILORED_DIR

# Valid document output formats
VALID_DOC_FORMATS = ("pdf", "docx")

log = logging.getLogger(__name__)


# ── Resume Parser ────────────────────────────────────────────────────────

def parse_resume(text: str) -> dict:
    """Parse a structured text resume into sections.

    Expects a format with header lines (name, title, location, contact)
    followed by ALL-CAPS section headers (SUMMARY, TECHNICAL SKILLS, etc.).

    Args:
        text: Full resume text.

    Returns:
        {"name": str, "title": str, "location": str, "contact": str, "sections": dict}
    """
    lines = [line.rstrip() for line in text.strip().split("\n")]

    # Header: first few lines before SUMMARY
    header_lines: list[str] = []
    body_start = 0
    for i, line in enumerate(lines):
        if line.strip().upper() == "SUMMARY":
            body_start = i
            break
        if line.strip():
            header_lines.append(line.strip())

    name = header_lines[0] if len(header_lines) > 0 else ""
    title = header_lines[1] if len(header_lines) > 1 else ""
    # The header may have 3 or 4 lines depending on whether location is included
    location = ""
    contact = ""
    if len(header_lines) > 3:
        location = header_lines[2]
        contact = header_lines[3]
    elif len(header_lines) > 2:
        # Could be location or contact -- check for email/phone indicators
        if "@" in header_lines[2] or "|" in header_lines[2]:
            contact = header_lines[2]
        else:
            location = header_lines[2]

    # Split body into sections by ALL-CAPS headers
    sections: dict[str, str] = {}
    current_section: str | None = None
    current_lines: list[str] = []

    for line in lines[body_start:]:
        stripped = line.strip()
        # Detect section headers (all caps, no leading dash/bullet, longer than 3 chars)
        if (
            stripped
            and stripped == stripped.upper()
            and not stripped.startswith("-")
            and len(stripped) > 3
            and not stripped.startswith("\u2022")
        ):
            if current_section:
                sections[current_section] = "\n".join(current_lines).strip()
            current_section = stripped
            current_lines = []
        else:
            current_lines.append(line)

    if current_section:
        sections[current_section] = "\n".join(current_lines).strip()

    return {
        "name": name,
        "title": title,
        "location": location,
        "contact": contact,
        "sections": sections,
    }


def parse_skills(text: str) -> list[tuple[str, str]]:
    """Parse skills section into (category, value) pairs.

    Args:
        text: The TECHNICAL SKILLS section text.

    Returns:
        List of (category_name, skills_string) tuples.
    """
    skills: list[tuple[str, str]] = []
    for line in text.strip().split("\n"):
        line = line.strip()
        if ":" in line:
            cat, val = line.split(":", 1)
            skills.append((cat.strip(), val.strip()))
    return skills


def parse_entries(text: str) -> list[dict]:
    """Parse experience/project entries from section text.

    Args:
        text: The EXPERIENCE or PROJECTS section text.

    Returns:
        List of {"title": str, "subtitle": str, "bullets": list[str]} dicts.
    """
    entries: list[dict] = []
    lines = text.strip().split("\n")
    current: dict | None = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("- ") or stripped.startswith("\u2022 "):
            if current:
                current["bullets"].append(stripped[2:].strip())
        elif current is None or (
            not stripped.startswith("-")
            and not stripped.startswith("\u2022")
            and len(current.get("bullets", [])) > 0
        ):
            # New entry
            if current:
                entries.append(current)
            current = {"title": stripped, "subtitle": "", "bullets": []}
        elif current and not current["subtitle"]:
            current["subtitle"] = stripped
        else:
            if current:
                current["bullets"].append(stripped)

    if current:
        entries.append(current)

    return entries


# ── HTML Template ────────────────────────────────────────────────────────

def build_html(resume: dict) -> str:
    """Build professional resume HTML from parsed data.

    Args:
        resume: Parsed resume dict from parse_resume().

    Returns:
        Complete HTML string ready for PDF rendering.
    """
    sections = resume["sections"]

    # Skills
    skills_html = ""
    if "TECHNICAL SKILLS" in sections:
        skills = parse_skills(sections["TECHNICAL SKILLS"])
        rows = ""
        for cat, val in skills:
            rows += f'<div class="skill-row"><span class="skill-cat">{cat}:</span> {val}</div>\n'
        skills_html = f'<div class="section"><div class="section-title">Technical Skills</div>{rows}</div>'

    # Experience
    exp_html = ""
    if "EXPERIENCE" in sections:
        entries = parse_entries(sections["EXPERIENCE"])
        items = ""
        for e in entries:
            bullets = "".join(f"<li>{b}</li>" for b in e["bullets"])
            subtitle = f'<div class="entry-subtitle">{e["subtitle"]}</div>' if e["subtitle"] else ""
            items += f'<div class="entry"><div class="entry-title">{e["title"]}</div>{subtitle}<ul>{bullets}</ul></div>'
        exp_html = f'<div class="section"><div class="section-title">Experience</div>{items}</div>'

    # Projects
    proj_html = ""
    if "PROJECTS" in sections:
        entries = parse_entries(sections["PROJECTS"])
        items = ""
        for e in entries:
            bullets = "".join(f"<li>{b}</li>" for b in e["bullets"])
            subtitle = f'<div class="entry-subtitle">{e["subtitle"]}</div>' if e["subtitle"] else ""
            items += f'<div class="entry"><div class="entry-title">{e["title"]}</div>{subtitle}<ul>{bullets}</ul></div>'
        proj_html = f'<div class="section"><div class="section-title">Projects</div>{items}</div>'

    # Education
    edu_html = ""
    if "EDUCATION" in sections:
        edu_text = sections["EDUCATION"].strip()
        edu_html = f'<div class="section"><div class="section-title">Education</div><div class="edu">{edu_text}</div></div>'

    # Summary
    summary_html = ""
    if "SUMMARY" in sections:
        summary_html = f'<div class="section"><div class="section-title">Summary</div><div class="summary">{sections["SUMMARY"].strip()}</div></div>'

    # Contact line parsing
    contact = resume["contact"]
    contact_parts = [p.strip() for p in contact.split("|")] if contact else []
    contact_html = " &nbsp;|&nbsp; ".join(contact_parts)

    # Location line (may be empty)
    location_html = f'<div class="location">{resume["location"]}</div>' if resume["location"] else ""

    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
@page {{
    size: letter;
    margin: 0.35in 0.5in;
}}
* {{
    margin: 0;
    padding: 0;
    box-sizing: border-box;
}}
body {{
    font-family: 'Calibri', 'Segoe UI', Arial, sans-serif;
    font-size: 10pt;
    line-height: 1.35;
    color: #1a1a1a;
}}
.header {{
    text-align: center;
    margin-bottom: 4px;
    padding-bottom: 4px;
    border-bottom: 1.5px solid #2a7ab5;
}}
.name {{
    font-size: 18pt;
    font-weight: 700;
    color: #1a3a5c;
    letter-spacing: 0.5px;
}}
.title {{
    font-size: 10.5pt;
    color: #3a6b8c;
    margin: 1px 0;
}}
.location {{
    font-size: 9pt;
    color: #555;
}}
.contact {{
    font-size: 9pt;
    color: #444;
    margin-top: 1px;
}}
.contact a {{
    color: #2c3e50;
    text-decoration: none;
}}
.section {{
    margin-top: 5px;
}}
.section-title {{
    font-size: 10pt;
    font-weight: 700;
    color: #1a3a5c;
    text-transform: uppercase;
    letter-spacing: 0.8px;
    border-bottom: 1.5px solid #2a7ab5;
    padding-bottom: 1px;
    margin-bottom: 3px;
}}
.summary {{
    font-size: 9.5pt;
    color: #333;
    line-height: 1.4;
}}
.skill-row {{
    font-size: 9.5pt;
    margin: 0;
    line-height: 1.35;
}}
.skill-cat {{
    font-weight: 600;
    color: #1a3a5c;
}}
.entry {{
    margin-bottom: 4px;
    break-inside: avoid;
}}
.entry-title {{
    font-weight: 600;
    font-size: 10pt;
    color: #1a3a5c;
}}
.entry-subtitle {{
    font-size: 9pt;
    color: #4a7a9b;
    font-style: italic;
    margin-bottom: 1px;
}}
ul {{
    margin-left: 14px;
    padding: 0;
}}
li {{
    font-size: 9.5pt;
    margin-bottom: 1px;
    line-height: 1.35;
}}
.edu {{
    font-size: 10pt;
}}
</style>
</head>
<body>
<div class="header">
    <div class="name">{resume['name']}</div>
    <div class="title">{resume['title']}</div>
    {location_html}
    <div class="contact">{contact_html}</div>
</div>
{summary_html}
{skills_html}
{exp_html}
{proj_html}
{edu_html}
</body>
</html>"""


# ── PDF Renderer ─────────────────────────────────────────────────────────

def render_pdf(html: str, output_path: str) -> None:
    """Render HTML to PDF using Playwright's headless Chromium.

    Args:
        html: Complete HTML string.
        output_path: Path to write the PDF file.
    """
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        page.set_content(html, wait_until="networkidle")
        page.pdf(
            path=output_path,
            format="Letter",
            margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
            print_background=True,
        )
        browser.close()


# ── DOCX Renderer ────────────────────────────────────────────────────

def render_docx(resume: dict, output_path: str, metadata: dict | None = None) -> None:
    """Render parsed resume data to a DOCX file using python-docx.

    Args:
        resume: Parsed resume dict from parse_resume().
        output_path: Path to write the DOCX file.
        metadata: Optional dict to populate the DOCX core_properties. Supported
            keys: 'title', 'subject', 'keywords' (str or list), 'description',
            'author', 'company', 'category', 'comments'.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins (match PDF: 0.35in top/bottom, 0.5in left/right)
    for section in doc.sections:
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.35)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.35

    # --- Header ---
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(resume["name"])
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    if resume["title"]:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(resume["title"])
        title_run.font.size = Pt(10.5)
        title_run.font.color.rgb = RGBColor(0x3A, 0x6B, 0x8C)

    if resume["location"]:
        loc_para = doc.add_paragraph()
        loc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        loc_run = loc_para.add_run(resume["location"])
        loc_run.font.size = Pt(9)
        loc_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if resume["contact"]:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(resume["contact"])
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    sections = resume["sections"]

    def _add_section_heading(title: str) -> None:
        """Add a styled section heading with bottom border."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(5)
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
        # Bottom border via XML (closest to the PDF blue line)
        from docx.oxml.ns import qn
        pPr = para._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): "2A7AB5",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    # --- Summary ---
    if "SUMMARY" in sections:
        _add_section_heading("Summary")
        p = doc.add_paragraph(sections["SUMMARY"].strip())
        p.runs[0].font.size = Pt(9.5)
        p.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- Technical Skills ---
    if "TECHNICAL SKILLS" in sections:
        _add_section_heading("Technical Skills")
        for cat, val in parse_skills(sections["TECHNICAL SKILLS"]):
            p = doc.add_paragraph()
            cat_run = p.add_run(f"{cat}: ")
            cat_run.bold = True
            cat_run.font.size = Pt(9.5)
            cat_run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
            val_run = p.add_run(val)
            val_run.font.size = Pt(9.5)

    # --- Experience (Jobscan §4: prefer "Work Experience" over bare "Experience") ---
    if "EXPERIENCE" in sections:
        _add_section_heading("Work Experience")
        for entry in parse_entries(sections["EXPERIENCE"]):
            tp = doc.add_paragraph()
            tp.paragraph_format.space_before = Pt(2)
            tr = tp.add_run(entry["title"])
            tr.bold = True
            tr.font.size = Pt(10)
            tr.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
            if entry["subtitle"]:
                sp = doc.add_paragraph()
                sr = sp.add_run(entry["subtitle"])
                sr.italic = True
                sr.font.size = Pt(9)
                sr.font.color.rgb = RGBColor(0x4A, 0x7A, 0x9B)
            for bullet in entry["bullets"]:
                bp = doc.add_paragraph(bullet, style="List Bullet")
                for run in bp.runs:
                    run.font.size = Pt(9.5)

    # --- Projects ---
    if "PROJECTS" in sections:
        _add_section_heading("Projects")
        for entry in parse_entries(sections["PROJECTS"]):
            tp = doc.add_paragraph()
            tp.paragraph_format.space_before = Pt(2)
            tr = tp.add_run(entry["title"])
            tr.bold = True
            tr.font.size = Pt(10)
            tr.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
            if entry["subtitle"]:
                sp = doc.add_paragraph()
                sr = sp.add_run(entry["subtitle"])
                sr.italic = True
                sr.font.size = Pt(9)
                sr.font.color.rgb = RGBColor(0x4A, 0x7A, 0x9B)
            for bullet in entry["bullets"]:
                bp = doc.add_paragraph(bullet, style="List Bullet")
                for run in bp.runs:
                    run.font.size = Pt(9.5)

    # --- Education ---
    if "EDUCATION" in sections:
        _add_section_heading("Education")
        p = doc.add_paragraph(sections["EDUCATION"].strip())
        p.runs[0].font.size = Pt(10)

    # Populate core_properties with metadata if provided.
    if metadata:
        cp = doc.core_properties
        if "title" in metadata and metadata["title"]:
            cp.title = str(metadata["title"])[:256]
        if "subject" in metadata and metadata["subject"]:
            cp.subject = str(metadata["subject"])[:256]
        if "author" in metadata and metadata["author"]:
            cp.author = str(metadata["author"])[:256]
        if "company" in metadata and metadata["company"]:
            cp.company = str(metadata["company"])[:256]
        if "category" in metadata and metadata["category"]:
            cp.category = str(metadata["category"])[:256]
        if "comments" in metadata and metadata["comments"]:
            cp.comments = str(metadata["comments"])[:2000]
        if "description" in metadata and metadata["description"]:
            # python-docx uses comments for long text; prefer comments for description
            # if both provided, comments has priority above.
            if not cp.comments:
                cp.comments = str(metadata["description"])[:2000]
        kw = metadata.get("keywords")
        if kw:
            if isinstance(kw, list):
                kw = ", ".join(str(k).strip() for k in kw if k)
            cp.keywords = str(kw)[:1024]

    doc.save(output_path)


# ── Public API ───────────────────────────────────────────────────────────

def convert_to_pdf(
    text_path: Path,
    output_path: Path | None = None,
    html_only: bool = False,
    doc_format: str = "docx",
    metadata: dict | None = None,
) -> Path:
    """Convert a text resume/cover letter to PDF or DOCX.

    Args:
        text_path: Path to the .txt file to convert.
        output_path: Optional override for the output path. Defaults to same
            name with the appropriate extension.
        html_only: If True, output HTML instead of PDF/DOCX.
        doc_format: Output format — "docx" (default) or "pdf".
        metadata: Optional dict to populate DOCX core_properties (ignored for PDF).
            Supported keys: 'title', 'subject', 'keywords', 'author', 'company',
            'category', 'comments', 'description'.

    Returns:
        Path to the generated file.
    """
    if doc_format not in VALID_DOC_FORMATS:
        raise ValueError(f"Invalid doc_format '{doc_format}'. Must be one of: {VALID_DOC_FORMATS}")

    text_path = Path(text_path)
    text = text_path.read_text(encoding="utf-8")
    resume = parse_resume(text)

    if html_only:
        html = build_html(resume)
        out = output_path or text_path.with_suffix(".html")
        out = Path(out)
        out.write_text(html, encoding="utf-8")
        log.info("HTML generated: %s", out)
        return out

    if doc_format == "docx":
        out = output_path or text_path.with_suffix(".docx")
        out = Path(out)
        render_docx(resume, str(out), metadata=metadata)
        log.info("DOCX generated: %s", out)
        return out

    # Default: PDF
    html = build_html(resume)
    out = output_path or text_path.with_suffix(".pdf")
    out = Path(out)
    render_pdf(html, str(out))
    log.info("PDF generated: %s", out)
    return out


def batch_convert(limit: int = 50, doc_format: str = "docx") -> int:
    """Convert .txt files in TAILORED_DIR that don't have corresponding output files.

    Scans for .txt files (excluding _JOB.txt and _REPORT.json), checks if a
    file with the target extension already exists, and converts any that are missing.

    Args:
        limit: Maximum number of files to convert.
        doc_format: Output format — "docx" (default) or "pdf".

    Returns:
        Number of files generated.
    """
    if doc_format not in VALID_DOC_FORMATS:
        raise ValueError(f"Invalid doc_format '{doc_format}'. Must be one of: {VALID_DOC_FORMATS}")

    ext = f".{doc_format}"

    if not TAILORED_DIR.exists():
        log.warning("Tailored directory does not exist: %s", TAILORED_DIR)
        return 0

    txt_files = sorted(TAILORED_DIR.glob("*.txt"))
    # Exclude _JOB.txt files from resume conversion
    # (they get their own conversion calls)
    candidates = [
        f for f in txt_files
        if not f.name.endswith("_JOB.txt")
    ]

    # Filter to those without a corresponding output file
    to_convert: list[Path] = []
    for f in candidates:
        out_path = f.with_suffix(ext)
        if not out_path.exists():
            to_convert.append(f)
        if len(to_convert) >= limit:
            break

    if not to_convert:
        log.debug("All text files already have %s files.", doc_format.upper())
        return 0

    log.info("Converting %d files to %s...", len(to_convert), doc_format.upper())
    converted = 0
    for f in to_convert:
        try:
            convert_to_pdf(f, doc_format=doc_format)
            converted += 1
        except Exception as e:
            log.error("Failed to convert %s: %s", f.name, e)

    log.info("Done: %d/%d %s files generated in %s", converted, len(to_convert), doc_format.upper(), TAILORED_DIR)
    return converted
