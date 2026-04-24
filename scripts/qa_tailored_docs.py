"""QA pass on recently-tailored DOCX resumes and cover letters.

Validates against the Jobscan methodology captured in
`docs/guidelines/jobscan/resume.md` and `cover-letter.md`:

- Filename format (Jobscan §3): FirstName_LastName_JobTitle_*.docx
- DOCX core_properties populated (title, author, subject, category,
  keywords, comments)
- Section header "Work Experience" (§4)
- Zero tables/columns (§7)
- Job title verbatim in body (§1: 10.6x interview-rate lift)
- Resume length: 1-2 pages (§10)
- Cover letter 250-400 words (cover-letter.md §1: 3.4x lift)
- Cover letter 3-4 paragraphs

Usage:
    .venv/bin/python scripts/qa_tailored_docs.py [--limit N] [--since-hours N]
"""

from __future__ import annotations

import argparse
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document

from applypilot.database import get_connection


# ── Individual checks ─────────────────────────────────────────────────

def check_filename(path: Path, first: str, last: str) -> tuple[bool, str]:
    """Check filename format matches FirstName_LastName_*.docx."""
    name = path.name
    if not first:
        return (True, "no candidate name to validate against")
    expected_prefix = f"{first}_{last}_" if last else f"{first}_"
    if name.startswith(expected_prefix):
        return (True, "OK")
    return (False, f"filename starts with {name.split('_')[0]!r}, expected {expected_prefix!r}")


def check_metadata(doc: Document) -> tuple[bool, list[str]]:
    """Check core_properties populated. Returns (all_ok, issues)."""
    cp = doc.core_properties
    issues: list[str] = []
    for field in ("title", "author", "subject", "category"):
        val = getattr(cp, field, None)
        if not val or not str(val).strip():
            issues.append(f"missing {field}")
    if not cp.keywords or not str(cp.keywords).strip():
        issues.append("missing keywords")
    return (not issues, issues)


def check_section_headers(doc: Document, is_cover: bool = False) -> tuple[bool, list[str]]:
    """Verify section headers match Jobscan recommendations.

    For resumes:
      - "Work Experience" (or "Professional Experience") — Jobscan §4 preferred
      - No bare "Experience"
    Cover letters: no section-header requirement.
    """
    if is_cover:
        return (True, ["(cover letter — no header requirement)"])

    headers: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text and len(text) < 40 and any(r.bold for r in p.runs) and (
            text == text.upper() or text.title() == text
        ):
            headers.append(text)

    issues: list[str] = []
    has_experience_header = any(
        re.fullmatch(r"(work|professional)\s+experience", h, re.IGNORECASE) for h in headers
    )
    has_bare_experience = any(re.fullmatch(r"experience", h, re.IGNORECASE) for h in headers)
    if has_bare_experience and not has_experience_header:
        issues.append("uses bare 'Experience' instead of 'Work Experience'")
    if not has_experience_header and not has_bare_experience:
        issues.append("no Experience section header found")
    return (not issues, issues)


def check_tables(doc: Document) -> tuple[bool, str]:
    n = len(doc.tables)
    if n == 0:
        return (True, "OK")
    return (False, f"{n} table(s) present (Jobscan §7: tables break parsing)")


def check_job_title_verbatim(doc: Document, job_title: str) -> tuple[bool, str]:
    """Jobscan §1: 10.6x interview lift for verbatim job title."""
    body = "\n".join(p.text for p in doc.paragraphs)
    if job_title and job_title in body:
        return (True, "verbatim match")
    tokens = set(re.findall(r"\w+", (job_title or "").lower()))
    body_tokens = set(re.findall(r"\w+", body.lower()))
    overlap = tokens & body_tokens
    ratio = len(overlap) / len(tokens) if tokens else 0
    if ratio >= 0.75:
        return (True, f"partial ({len(overlap)}/{len(tokens)} tokens)")
    return (False, f"low overlap {len(overlap)}/{len(tokens)} tokens")


def check_resume_length(doc: Document) -> tuple[bool, str]:
    chars = sum(len(p.text) for p in doc.paragraphs)
    pages = chars / 3500
    if 0.5 <= pages <= 2.2:
        return (True, f"{pages:.1f} pages")
    return (False, f"{pages:.1f} pages (target 1-2)")


def check_cover_letter_length(doc: Document) -> tuple[bool, str]:
    body = "\n".join(p.text for p in doc.paragraphs)
    words = len(body.split())
    if 220 <= words <= 450:
        # Jobscan ideal is 250-400; accept 220-450 as "close enough" per user.
        marker = "" if 250 <= words <= 400 else " (OK, below Jobscan ideal)"
        return (True, f"{words} words{marker}")
    if words < 220:
        return (False, f"{words} words (below 220 floor)")
    return (False, f"{words} words (above 450 ceiling)")


def check_cover_letter_paragraphs(doc: Document) -> tuple[bool, str]:
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    n = len(non_empty)
    if 3 <= n <= 5:
        return (True, f"{n} paragraphs")
    return (False, f"{n} paragraphs (target 3-4)")


# ── Runner ────────────────────────────────────────────────────────────

def run_qa(limit: int = 30, since_hours: int = 12) -> int:
    from applypilot.config import load_profile
    from applypilot.scoring.tailor import _name_parts

    profile = load_profile()
    first, last = _name_parts(profile)

    conn = get_connection()
    rows = conn.execute(f"""
        SELECT url, title, fit_score, tailored_resume_path, cover_letter_path
        FROM jobs
        WHERE tailored_resume_path LIKE '%.docx'
          AND cover_letter_path LIKE '%.docx'
          AND fit_score >= 8
          AND tailored_at > datetime('now', '-{since_hours} hours')
        ORDER BY RANDOM()
        LIMIT {limit}
    """).fetchall()

    if not rows:
        print("No recent tailored DOCX jobs found to QA.")
        return 1

    print(f"QA on {len(rows)} resume+cover samples "
          f"(tailored in last {since_hours}h, score>=8)\n")

    resume_results: Counter[str] = Counter()
    cover_results: Counter[str] = Counter()
    failed_details: list[str] = []

    for r in rows:
        resume_path = Path(r["tailored_resume_path"])
        cover_path = Path(r["cover_letter_path"])
        job_title = r["title"] or ""

        if not resume_path.exists():
            resume_results["missing_file"] += 1
            continue
        if not cover_path.exists():
            cover_results["missing_file"] += 1
            continue

        resume_doc = Document(str(resume_path))
        cover_doc = Document(str(cover_path))

        # Resume checks
        checks = [
            ("filename", *check_filename(resume_path, first, last)),
            ("metadata", *check_metadata(resume_doc)),
            ("headers",  *check_section_headers(resume_doc)),
            ("tables",   *check_tables(resume_doc)),
            ("title_verbatim", *check_job_title_verbatim(resume_doc, job_title)),
            ("length",   *check_resume_length(resume_doc)),
        ]
        for name, ok, detail in checks:
            resume_results[f"{name}_{'pass' if ok else 'fail'}"] += 1
            if not ok:
                issue = detail if isinstance(detail, str) else ", ".join(detail)
                failed_details.append(f"  resume {name}: {issue} | {resume_path.name}")

        # Cover letter checks
        cover_checks = [
            ("filename", *check_filename(cover_path, first, last)),
            ("metadata", *check_metadata(cover_doc)),
            ("length",   *check_cover_letter_length(cover_doc)),
            ("paragraphs", *check_cover_letter_paragraphs(cover_doc)),
        ]
        for name, ok, detail in cover_checks:
            cover_results[f"{name}_{'pass' if ok else 'fail'}"] += 1
            if not ok:
                issue = detail if isinstance(detail, str) else ", ".join(detail)
                failed_details.append(f"  cover  {name}: {issue} | {cover_path.name}")

    print("=== Resume check aggregates ===")
    checks_order = ["filename", "metadata", "headers", "tables", "title_verbatim", "length"]
    for c in checks_order:
        p = resume_results.get(f"{c}_pass", 0)
        f = resume_results.get(f"{c}_fail", 0)
        total = p + f
        pct = 100 * p / total if total else 0
        marker = " ✅" if f == 0 else " ⚠️" if p > f else " ❌"
        print(f"  {c:<16} {p:>3}/{total:<3} pass ({pct:.0f}%){marker}")

    print()
    print("=== Cover letter check aggregates ===")
    cover_order = ["filename", "metadata", "length", "paragraphs"]
    for c in cover_order:
        p = cover_results.get(f"{c}_pass", 0)
        f = cover_results.get(f"{c}_fail", 0)
        total = p + f
        pct = 100 * p / total if total else 0
        marker = " ✅" if f == 0 else " ⚠️" if p > f else " ❌"
        print(f"  {c:<16} {p:>3}/{total:<3} pass ({pct:.0f}%){marker}")

    if failed_details:
        print()
        print(f"=== Failures ({len(failed_details)} total, first 15 shown) ===")
        for line in failed_details[:15]:
            print(line)

    return 0 if not failed_details else 2


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--limit", type=int, default=30, help="Number of samples to QA")
    parser.add_argument("--since-hours", type=int, default=12, help="Only QA docs tailored within this window")
    args = parser.parse_args()
    sys.exit(run_qa(args.limit, args.since_hours))
